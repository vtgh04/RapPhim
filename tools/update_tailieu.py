"""
Script to rewrite TaiLieu.docx with the updated specification.
Reads existing document structure and rewrites relevant sections
to match the new entity definitions, removing Service, InvoiceDetail, etc.

Run: python tools/update_tailieu.py
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import copy
import os

INPUT_PATH = r'c:\Users\ADMIN\Desktop\CinemaProject\RapPhim\docs\TaiLieu.docx'
OUTPUT_PATH = r'c:\Users\ADMIN\Desktop\CinemaProject\RapPhim\docs\TaiLieu.docx'
BACKUP_PATH = r'c:\Users\ADMIN\Desktop\CinemaProject\RapPhim\docs\TaiLieu_backup.docx'

# ─── Step 1: Read existing document to understand structure ───
print("Reading existing document...")
doc = Document(INPUT_PATH)

# Dump current content for reference
print("\n=== CURRENT DOCUMENT STRUCTURE ===")
for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        print(f"  [{i:3d}][{p.style.name:20s}] {p.text[:100]}")

print(f"\n  Tables found: {len(doc.tables)}")
for ti, table in enumerate(doc.tables):
    print(f"  Table {ti}: {len(table.rows)} rows x {len(table.columns)} cols")
    if table.rows:
        first_row = [cell.text[:30] for cell in table.rows[0].cells]
        print(f"    Header: {' | '.join(first_row)}")


# ─── Step 2: Save backup ───
print(f"\nSaving backup to {BACKUP_PATH}...")
doc.save(BACKUP_PATH)


# ─── Step 3: Helper functions ───
def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading_elm = cell._element.get_or_add_tcPr()
    shading = shading_elm.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    shading_elm.append(shading)

def add_formatted_table(doc, headers, rows, col_widths=None):
    """Add a nicely formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    header_row = table.rows[0]
    for i, header_text in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header_text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, 'D9E2F3')  # Light blue header
    
    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'
    
    # Set column widths if provided
    if col_widths:
        for row in table.rows:
            for ci, width in enumerate(col_widths):
                row.cells[ci].width = Cm(width)
    
    return table

def add_heading_styled(doc, text, level=1):
    """Add a heading with consistent styling."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
    return h

def add_paragraph_styled(doc, text, bold=False, italic=False, size=12):
    """Add a paragraph with consistent styling."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p

def add_bullet(doc, text, bold_prefix=None):
    """Add a bullet point."""
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    else:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    return p


# ─── Step 4: Rewrite the document ───
print("\nRewriting document with updated specification...")

new_doc = Document()

# Set default font
style = new_doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# ============================================================
# TITLE
# ============================================================
title = new_doc.add_heading('TÀI LIỆU ĐẶC TẢ HỆ THỐNG', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = new_doc.add_heading('Phần mềm Quản lý Rạp Phim (RapPhim)', level=1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

new_doc.add_paragraph('')

# ============================================================
# 1. TỔNG QUAN HỆ THỐNG
# ============================================================
add_heading_styled(new_doc, '1. Tổng quan hệ thống', level=1)

add_paragraph_styled(new_doc, 
    'Hệ thống quản lý rạp phim (RapPhim) là ứng dụng desktop Java Swing hỗ trợ '
    'nhân viên bán vé tại quầy, quản lý suất chiếu, phòng chiếu, phim và các chương '
    'trình giảm giá. Hệ thống phân quyền theo vai trò: Quản lý (MANAGER) và Nhân viên (STAFF).')

add_heading_styled(new_doc, '1.1. Phạm vi hệ thống', level=2)
features = [
    'Quản lý phim (thêm, sửa, ẩn/hiện)',
    'Quản lý phòng chiếu và ghế ngồi',
    'Quản lý suất chiếu (lịch chiếu, tránh trùng lịch)',
    'Bán vé tại quầy (chọn ghế, giữ chỗ, thanh toán, in vé)',
    'Quản lý hóa đơn và vé',
    'Quản lý chương trình giảm giá',
    'Quản lý nhân viên và phân quyền đăng nhập',
    'Thống kê doanh thu',
]
for f in features:
    add_bullet(new_doc, f)

add_heading_styled(new_doc, '1.2. Quy ước đặt tên', level=2)
add_paragraph_styled(new_doc, 
    'Ở mức OOP (Java): tên class dùng số ít (Employee, Movie, Ticket, ...). '
    'Ở mức Database: tên bảng dùng số nhiều (employees, movies, tickets, ...). '
    'Khóa chính dùng dạng camelCase trong Java (employeeId) và snake_case trong DB (employee_id).')

# ============================================================
# 2. ĐẶC TẢ THỰC THỂ (ENTITY SPECIFICATION)
# ============================================================
add_heading_styled(new_doc, '2. Đặc tả thực thể', level=1)

add_paragraph_styled(new_doc, 
    'Hệ thống gồm 9 lớp thực thể chính. Đã loại bỏ class Service và InvoiceDetail '
    'vì scope hiện tại chỉ bán vé, mô hình Invoice → Ticket là đủ, tránh dư thừa '
    'và mâu thuẫn quantity với ticket_id.')

# ─── 2.1 Employee ───
add_heading_styled(new_doc, '2.1. Employee (Nhân viên)', level=2)
add_paragraph_styled(new_doc, 'Vai trò: Quản lý đăng nhập và phân quyền người dùng hệ thống.', italic=True)

add_formatted_table(new_doc,
    ['Thuộc tính', 'Kiểu dữ liệu', 'Ràng buộc', 'Mô tả'],
    [
        ['employeeId', 'String (PK)', 'NOT NULL, AUTO', 'Mã nhân viên, khóa chính'],
        ['fullName', 'String', 'NOT NULL', 'Họ tên đầy đủ'],
        ['username', 'String', 'NOT NULL, UNIQUE', 'Tên đăng nhập, duy nhất'],
        ['passwordHash', 'String', 'NOT NULL', 'Mật khẩu đã hash (BCrypt)'],
        ['role', 'Enum', 'NOT NULL', 'MANAGER | STAFF'],
        ['status', 'Enum', 'NOT NULL, DEFAULT ACTIVE', 'ACTIVE | INACTIVE'],
        ['phone', 'String', 'Nullable', 'Số điện thoại'],
        ['email', 'String', 'UNIQUE (nullable)', 'Email, nên duy nhất'],
    ],
    col_widths=[3.5, 3, 4, 5.5]
)
new_doc.add_paragraph('')
add_paragraph_styled(new_doc, 'Ràng buộc nghiệp vụ:', bold=True)
add_bullet(new_doc, 'username là duy nhất trong hệ thống')
add_bullet(new_doc, 'email nên duy nhất (nullable)')
add_bullet(new_doc, 'Chỉ tài khoản có status = ACTIVE mới được đăng nhập')
add_bullet(new_doc, 'Logic điều hướng: role = MANAGER → Admin page, role = STAFF → Staff page')

# ─── 2.2 Movie ───
add_heading_styled(new_doc, '2.2. Movie (Phim)', level=2)
add_paragraph_styled(new_doc, 'Vai trò: Lưu thông tin phim chiếu tại rạp.', italic=True)

add_formatted_table(new_doc,
    ['Thuộc tính', 'Kiểu dữ liệu', 'Ràng buộc', 'Mô tả'],
    [
        ['movieId', 'String (PK)', 'NOT NULL, AUTO', 'Mã phim, khóa chính'],
        ['title', 'String', 'NOT NULL', 'Tên phim'],
        ['genre', 'String', 'NOT NULL', 'Thể loại phim'],
        ['durationMins', 'int', 'NOT NULL, > 0', 'Thời lượng (phút)'],
        ['language', 'String', 'Nullable', 'Ngôn ngữ / phụ đề'],
        ['releaseDate', 'LocalDate', 'Nullable', 'Ngày phát hành'],
        ['status', 'Enum', 'NOT NULL, DEFAULT ACTIVE', 'ACTIVE | INACTIVE'],
        ['description', 'String (TEXT)', 'Nullable', 'Mô tả / tóm tắt phim'],
        ['posterUrl', 'String', 'Nullable', 'URL ảnh poster'],
    ],
    col_widths=[3.5, 3, 4, 5.5]
)
new_doc.add_paragraph('')
add_paragraph_styled(new_doc, 'Ràng buộc nghiệp vụ:', bold=True)
add_bullet(new_doc, 'durationMins phải > 0')
add_bullet(new_doc, 'Không xóa cứng (DELETE) phim đã phát sinh suất chiếu hoặc vé → chỉ đặt status = INACTIVE')

# ─── 2.3 CinemaHall ───
add_heading_styled(new_doc, '2.3. CinemaHall (Phòng chiếu)', level=2)
add_paragraph_styled(new_doc, 'Vai trò: Lưu thông tin phòng chiếu phim.', italic=True)

add_formatted_table(new_doc,
    ['Thuộc tính', 'Kiểu dữ liệu', 'Ràng buộc', 'Mô tả'],
    [
        ['hallId', 'String (PK)', 'NOT NULL, AUTO', 'Mã phòng chiếu, khóa chính'],
        ['name', 'String', 'NOT NULL, UNIQUE', 'Tên phòng (vd: Phòng 1, Hall A)'],
        ['hallType', 'String', 'Nullable', 'Loại phòng (2D, 3D, IMAX, ...)'],
        ['totalRows', 'int', 'NOT NULL, > 0', 'Tổng số hàng ghế'],
        ['totalCols', 'int', 'NOT NULL, > 0', 'Tổng số cột ghế'],
        ['isActive', 'boolean', 'NOT NULL, DEFAULT true', 'Phòng còn hoạt động'],
    ],
    col_widths=[3.5, 3, 4, 5.5]
)
new_doc.add_paragraph('')
add_paragraph_styled(new_doc, 'Ràng buộc nghiệp vụ:', bold=True)
add_bullet(new_doc, 'totalRows > 0, totalCols > 0')
add_bullet(new_doc, 'name không được trùng (UNIQUE)')

# ─── 2.4 Seat ───
add_heading_styled(new_doc, '2.4. Seat (Ghế ngồi)', level=2)
add_paragraph_styled(new_doc, 'Vai trò: Ghế vật lý cố định trong phòng chiếu.', italic=True)

add_formatted_table(new_doc,
    ['Thuộc tính', 'Kiểu dữ liệu', 'Ràng buộc', 'Mô tả'],
    [
        ['seatId', 'String (PK)', 'NOT NULL, AUTO', 'Mã ghế, khóa chính'],
        ['hallId', 'String (FK)', 'NOT NULL', 'FK → CinemaHall.hallId'],
        ['rowChar', 'char', 'NOT NULL', 'Ký tự hàng (A, B, C, ...)'],
        ['colNumber', 'int', 'NOT NULL', 'Số cột (1, 2, 3, ...)'],
        ['seatType', 'Enum', 'NOT NULL, DEFAULT REGULAR', 'REGULAR | VIP | ACCESSIBLE'],
    ],
    col_widths=[3.5, 3, 4, 5.5]
)
new_doc.add_paragraph('')
add_paragraph_styled(new_doc, 'Ràng buộc nghiệp vụ:', bold=True)
add_bullet(new_doc, 'Duy nhất theo bộ (hallId, rowChar, colNumber) — UNIQUE constraint')
add_bullet(new_doc, 'Enum seatType: REGULAR (thường), VIP (cao cấp), ACCESSIBLE (hỗ trợ)')

# ─── 2.5 Showtime ───
add_heading_styled(new_doc, '2.5. Showtime (Suất chiếu)', level=2)
add_paragraph_styled(new_doc, 'Vai trò: Một lần chiếu cụ thể của một phim trong một phòng.', italic=True)

add_formatted_table(new_doc,
    ['Thuộc tính', 'Kiểu dữ liệu', 'Ràng buộc', 'Mô tả'],
    [
        ['showtimeId', 'String (PK)', 'NOT NULL, AUTO', 'Mã suất chiếu, khóa chính'],
        ['movieId', 'String (FK)', 'NOT NULL', 'FK → Movie.movieId'],
        ['hallId', 'String (FK)', 'NOT NULL', 'FK → CinemaHall.hallId'],
        ['startTime', 'LocalDateTime', 'NOT NULL', 'Thời gian bắt đầu chiếu'],
        ['endTime', 'LocalDateTime', 'NOT NULL', 'Thời gian kết thúc (tính cả dọn phòng)'],
        ['basePrice', 'double', 'NOT NULL, >= 0', 'Giá vé cơ bản của suất chiếu'],
        ['status', 'Enum', 'NOT NULL, DEFAULT SCHEDULED', 'SCHEDULED | ONGOING | COMPLETED | CANCELLED'],
    ],
    col_widths=[3.5, 3, 4, 5.5]
)
new_doc.add_paragraph('')
add_paragraph_styled(new_doc, 'Ràng buộc nghiệp vụ:', bold=True)
add_bullet(new_doc, 'endTime > startTime')
add_bullet(new_doc, 'Logic thời gian: endTime = startTime + durationMins + 15 phút dọn phòng')
add_bullet(new_doc, 'Không được chồng lịch trong cùng hallId (kiểm tra overlap trước khi tạo)')

# ─── 2.6 ShowSeat ───
add_heading_styled(new_doc, '2.6. ShowSeat (Ghế trong suất chiếu)', level=2)
add_paragraph_styled(new_doc, 'Vai trò: Trạng thái của một ghế trong một suất chiếu cụ thể.', italic=True)

add_formatted_table(new_doc,
    ['Thuộc tính', 'Kiểu dữ liệu', 'Ràng buộc', 'Mô tả'],
    [
        ['showSeatId', 'String (PK)', 'NOT NULL, AUTO', 'Mã ghế-suất chiếu, khóa chính'],
        ['showtimeId', 'String (FK)', 'NOT NULL', 'FK → Showtime.showtimeId'],
        ['seatId', 'String (FK)', 'NOT NULL', 'FK → Seat.seatId'],
        ['price', 'double', 'NOT NULL, >= 0', 'Giá vé thực tế cho ghế này'],
        ['status', 'Enum', 'NOT NULL, DEFAULT AVAILABLE', 'AVAILABLE | HELD | BOOKED'],
        ['heldUntil', 'LocalDateTime', 'Nullable', 'Thời hạn giữ chỗ (quan trọng)'],
    ],
    col_widths=[3.5, 3, 4, 5.5]
)
new_doc.add_paragraph('')
add_paragraph_styled(new_doc, 'Ràng buộc nghiệp vụ:', bold=True)
add_bullet(new_doc, 'Duy nhất theo bộ (showtimeId, seatId) — UNIQUE constraint')
add_bullet(new_doc, 'Enum status: AVAILABLE (trống), HELD (đang giữ chỗ), BOOKED (đã đặt)')
add_bullet(new_doc, 'heldUntil rất quan trọng để ghế giữ chỗ không bị kẹt vô hạn — cần job/schedule để giải phóng ghế hết hạn')

# ─── 2.7 Invoice ───
add_heading_styled(new_doc, '2.7. Invoice (Hóa đơn)', level=2)
add_paragraph_styled(new_doc, 'Vai trò: Hóa đơn thanh toán khi bán vé.', italic=True)

add_formatted_table(new_doc,
    ['Thuộc tính', 'Kiểu dữ liệu', 'Ràng buộc', 'Mô tả'],
    [
        ['invoiceId', 'String (PK)', 'NOT NULL, AUTO', 'Mã hóa đơn, khóa chính'],
        ['employeeId', 'String (FK)', 'NOT NULL', 'FK → Employee.employeeId (người tạo)'],
        ['createdAt', 'LocalDateTime', 'NOT NULL, DEFAULT NOW', 'Thời gian tạo hóa đơn'],
        ['totalAmount', 'double', 'NOT NULL, >= 0', 'Tổng tiền (tính từ các Ticket)'],
        ['paymentMethod', 'Enum', 'NOT NULL', 'CASH | CARD | TRANSFER'],
        ['status', 'Enum', 'NOT NULL, DEFAULT PENDING', 'PENDING | CONFIRMED | CANCELLED'],
        ['note', 'String (TEXT)', 'Nullable', 'Ghi chú thêm'],
    ],
    col_widths=[3.5, 3, 4, 5.5]
)
new_doc.add_paragraph('')
add_paragraph_styled(new_doc, 'Ràng buộc nghiệp vụ:', bold=True)
add_bullet(new_doc, 'Enum paymentMethod: CASH (tiền mặt), CARD (thẻ), TRANSFER (chuyển khoản)')
add_bullet(new_doc, 'Enum status: PENDING (chờ), CONFIRMED (xác nhận), CANCELLED (đã hủy)')
add_bullet(new_doc, 'totalAmount = SUM(finalPrice) của tất cả Ticket thuộc Invoice')
add_bullet(new_doc, 'Đã bỏ InvoiceDetail — Invoice liên kết trực tiếp tới Ticket (1-N)')

# ─── 2.8 Ticket ───
add_heading_styled(new_doc, '2.8. Ticket (Vé)', level=2)
add_paragraph_styled(new_doc, 'Vai trò: Vé đã bán, gắn với một ghế cụ thể trong một suất chiếu.', italic=True)

add_formatted_table(new_doc,
    ['Thuộc tính', 'Kiểu dữ liệu', 'Ràng buộc', 'Mô tả'],
    [
        ['ticketId', 'String (PK)', 'NOT NULL, AUTO', 'Mã vé, khóa chính'],
        ['invoiceId', 'String (FK)', 'NOT NULL', 'FK → Invoice.invoiceId'],
        ['showSeatId', 'String (FK)', 'NOT NULL', 'FK → ShowSeat.showSeatId'],
        ['discountId', 'String (FK)', 'Nullable', 'FK → Discount.discountId (nếu áp dụng)'],
        ['barcode', 'String', 'UNIQUE', 'Mã vạch duy nhất để quét vé'],
        ['originalPrice', 'double', 'NOT NULL, >= 0', 'Giá gốc'],
        ['discountAmount', 'double', 'NOT NULL, DEFAULT 0', 'Số tiền được giảm'],
        ['finalPrice', 'double', 'NOT NULL, >= 0', 'Giá cuối = originalPrice - discountAmount'],
        ['issuedAt', 'LocalDateTime', 'NOT NULL, DEFAULT NOW', 'Thời gian xuất vé'],
        ['status', 'Enum', 'NOT NULL, DEFAULT VALID', 'VALID | USED | CANCELLED'],
    ],
    col_widths=[3.5, 3, 4, 5.5]
)
new_doc.add_paragraph('')
add_paragraph_styled(new_doc, 'Ràng buộc nghiệp vụ:', bold=True)
add_bullet(new_doc, 'Một showSeat chỉ có tối đa một vé hợp lệ (status != CANCELLED)')
add_bullet(new_doc, 'finalPrice = originalPrice - discountAmount')
add_bullet(new_doc, 'Enum status: VALID (hợp lệ), USED (đã sử dụng), CANCELLED (đã hủy)')
add_bullet(new_doc, 'discountId nullable — chỉ có giá trị khi vé áp dụng chương trình giảm giá')

# ─── 2.9 Discount ───
add_heading_styled(new_doc, '2.9. Discount (Giảm giá)', level=2)
add_paragraph_styled(new_doc, 'Vai trò: Định nghĩa chương trình giảm giá.', italic=True)

add_formatted_table(new_doc,
    ['Thuộc tính', 'Kiểu dữ liệu', 'Ràng buộc', 'Mô tả'],
    [
        ['discountId', 'String (PK)', 'NOT NULL, AUTO', 'Mã giảm giá, khóa chính'],
        ['discountName', 'String', 'NOT NULL', 'Tên chương trình giảm giá'],
        ['discountType', 'Enum', 'NOT NULL', 'HOLIDAY | GROUP | SPECIAL'],
        ['discountRate', 'double', 'NOT NULL, 0-1', 'Tỷ lệ giảm (vd: 0.1 = 10%)'],
        ['validFrom', 'LocalDate', 'NOT NULL', 'Ngày bắt đầu hiệu lực'],
        ['validTo', 'LocalDate', 'NOT NULL', 'Ngày kết thúc hiệu lực'],
        ['minTicketQuantity', 'int', 'NOT NULL, DEFAULT 1', 'Số vé tối thiểu để áp dụng'],
        ['isActive', 'boolean', 'NOT NULL, DEFAULT true', 'Chương trình đang hoạt động'],
        ['description', 'String (TEXT)', 'Nullable', 'Mô tả chi tiết'],
    ],
    col_widths=[3.5, 3, 4, 5.5]
)
new_doc.add_paragraph('')
add_paragraph_styled(new_doc, 'Ràng buộc nghiệp vụ:', bold=True)
add_bullet(new_doc, 'Enum discountType: HOLIDAY (ngày lễ), GROUP (mua nhóm), SPECIAL (đặc biệt)')
add_bullet(new_doc, 'validTo >= validFrom')
add_bullet(new_doc, 'Giảm giá chỉ áp dụng khi isActive = true và ngày hiện tại trong [validFrom, validTo]')

# ============================================================
# 3. QUAN HỆ GIỮA CÁC THỰC THỂ
# ============================================================
add_heading_styled(new_doc, '3. Quan hệ giữa các thực thể', level=1)

add_formatted_table(new_doc,
    ['Thực thể A', 'Quan hệ', 'Thực thể B', 'Ghi chú'],
    [
        ['Movie', '1 — N', 'Showtime', 'Một phim có nhiều suất chiếu'],
        ['CinemaHall', '1 — N', 'Seat', 'Một phòng có nhiều ghế'],
        ['CinemaHall', '1 — N', 'Showtime', 'Một phòng có nhiều suất chiếu'],
        ['Showtime', '1 — N', 'ShowSeat', 'Một suất chiếu tạo ra N bản ghi ShowSeat'],
        ['Seat', '1 — N', 'ShowSeat', 'Một ghế xuất hiện trong nhiều suất chiếu'],
        ['Employee', '1 — N', 'Invoice', 'Một nhân viên tạo nhiều hóa đơn'],
        ['Invoice', '1 — N', 'Ticket', 'Một hóa đơn chứa nhiều vé (không qua InvoiceDetail)'],
        ['Discount', '1 — N', 'Ticket', 'Một chương trình giảm giá áp dụng cho nhiều vé (optional)'],
        ['ShowSeat', '1 — 0..1', 'Ticket', 'Một ghế suất chiếu có tối đa 1 vé hợp lệ'],
    ],
    col_widths=[3, 2, 3, 8]
)

# ============================================================
# 4. NHỮNG GÌ ĐÃ BỎ (SO VỚI BẢN TRƯỚC)
# ============================================================
add_heading_styled(new_doc, '4. Những thay đổi so với bản trước', level=1)

add_paragraph_styled(new_doc, 'Các thành phần đã loại bỏ:', bold=True)
add_bullet(new_doc, 'Bỏ class Service — scope hiện tại chỉ bán vé, không bán dịch vụ kèm theo')
add_bullet(new_doc, 'Bỏ class InvoiceDetail — Invoice liên kết trực tiếp tới Ticket (1-N), tránh dư thừa và mâu thuẫn quantity/ticket_id')
add_bullet(new_doc, 'Bỏ các thuộc tính itemType, serviceId, quantity kiểu line-item hỗn hợp trong hóa đơn')
add_bullet(new_doc, 'Bỏ logic "ghế ACCESSIBLE thì tự động giảm giá người khuyết tật" — sai nghiệp vụ (loại ghế ≠ đối tượng giảm giá)')

# ============================================================
# 5. LOGIC ĐĂNG NHẬP
# ============================================================
add_heading_styled(new_doc, '5. Logic đăng nhập', level=1)

add_paragraph_styled(new_doc, 'Quy trình xác thực và điều hướng:', bold=True)
numbered = [
    '1. Nhân viên nhập username và password.',
    '2. Hệ thống kiểm tra username tồn tại → nếu không: báo lỗi.',
    '3. Hệ thống kiểm tra status == ACTIVE → nếu không: báo "Tài khoản đã bị vô hiệu hóa".',
    '4. Hệ thống xác minh password (so sánh hash) → nếu sai: báo lỗi mật khẩu.',
    '5. Đăng nhập thành công → điều hướng theo role:',
]
for item in numbered:
    add_paragraph_styled(new_doc, item)

add_bullet(new_doc, 'role = MANAGER → mở Admin page (quản lý phim, phòng, nhân viên, thống kê, ...)')
add_bullet(new_doc, 'role = STAFF → mở Staff page (bán vé, xem lịch chiếu, ...)')


# ============================================================
# 6. THIẾT KẾ DATABASE (SQL)
# ============================================================
add_heading_styled(new_doc, '6. Thiết kế Database', level=1)

add_paragraph_styled(new_doc, 
    'Tên bảng dùng số nhiều (employees, movies, ...). '
    'Khóa chính dùng snake_case. Kiểu dữ liệu theo SQL Server / MySQL.')

# employees table
add_heading_styled(new_doc, '6.1. Bảng employees', level=2)
add_formatted_table(new_doc,
    ['Cột', 'Kiểu', 'Ràng buộc', 'Mô tả'],
    [
        ['employee_id', 'VARCHAR(20)', 'PK', 'Mã nhân viên'],
        ['full_name', 'NVARCHAR(100)', 'NOT NULL', 'Họ tên'],
        ['username', 'VARCHAR(50)', 'NOT NULL, UNIQUE', 'Tên đăng nhập'],
        ['password_hash', 'VARCHAR(255)', 'NOT NULL', 'Mật khẩu hash'],
        ['role', 'VARCHAR(20)', "NOT NULL, CHECK IN ('MANAGER','STAFF')", 'Vai trò'],
        ['status', 'VARCHAR(20)', "NOT NULL, DEFAULT 'ACTIVE'", 'ACTIVE | INACTIVE'],
        ['phone', 'VARCHAR(20)', 'NULL', 'SĐT'],
        ['email', 'VARCHAR(100)', 'UNIQUE, NULL', 'Email'],
    ]
)

# movies table
add_heading_styled(new_doc, '6.2. Bảng movies', level=2)
add_formatted_table(new_doc,
    ['Cột', 'Kiểu', 'Ràng buộc', 'Mô tả'],
    [
        ['movie_id', 'VARCHAR(20)', 'PK', 'Mã phim'],
        ['title', 'NVARCHAR(200)', 'NOT NULL', 'Tên phim'],
        ['genre', 'NVARCHAR(100)', 'NOT NULL', 'Thể loại'],
        ['duration_mins', 'INT', 'NOT NULL, CHECK > 0', 'Thời lượng (phút)'],
        ['language', 'NVARCHAR(50)', 'NULL', 'Ngôn ngữ'],
        ['release_date', 'DATE', 'NULL', 'Ngày phát hành'],
        ['status', 'VARCHAR(20)', "NOT NULL, DEFAULT 'ACTIVE'", 'ACTIVE | INACTIVE'],
        ['description', 'NTEXT', 'NULL', 'Mô tả phim'],
        ['poster_url', 'VARCHAR(500)', 'NULL', 'URL poster'],
    ]
)

# cinema_halls table
add_heading_styled(new_doc, '6.3. Bảng cinema_halls', level=2)
add_formatted_table(new_doc,
    ['Cột', 'Kiểu', 'Ràng buộc', 'Mô tả'],
    [
        ['hall_id', 'VARCHAR(20)', 'PK', 'Mã phòng'],
        ['name', 'NVARCHAR(100)', 'NOT NULL, UNIQUE', 'Tên phòng'],
        ['hall_type', 'VARCHAR(20)', 'NULL', 'Loại phòng (2D, 3D, IMAX)'],
        ['total_rows', 'INT', 'NOT NULL, CHECK > 0', 'Tổng hàng'],
        ['total_cols', 'INT', 'NOT NULL, CHECK > 0', 'Tổng cột'],
        ['is_active', 'BIT', 'NOT NULL, DEFAULT 1', 'Đang hoạt động'],
    ]
)

# seats table
add_heading_styled(new_doc, '6.4. Bảng seats', level=2)
add_formatted_table(new_doc,
    ['Cột', 'Kiểu', 'Ràng buộc', 'Mô tả'],
    [
        ['seat_id', 'VARCHAR(20)', 'PK', 'Mã ghế'],
        ['hall_id', 'VARCHAR(20)', 'FK → cinema_halls', 'Phòng chứa ghế'],
        ['row_char', 'CHAR(1)', 'NOT NULL', 'Ký tự hàng (A-Z)'],
        ['col_number', 'INT', 'NOT NULL', 'Số cột'],
        ['seat_type', 'VARCHAR(20)', "NOT NULL, DEFAULT 'REGULAR'", 'REGULAR | VIP | ACCESSIBLE'],
    ]
)
add_paragraph_styled(new_doc, 'UNIQUE constraint: (hall_id, row_char, col_number)', italic=True)

# showtimes table
add_heading_styled(new_doc, '6.5. Bảng showtimes', level=2)
add_formatted_table(new_doc,
    ['Cột', 'Kiểu', 'Ràng buộc', 'Mô tả'],
    [
        ['showtime_id', 'VARCHAR(20)', 'PK', 'Mã suất chiếu'],
        ['movie_id', 'VARCHAR(20)', 'FK → movies', 'Phim chiếu'],
        ['hall_id', 'VARCHAR(20)', 'FK → cinema_halls', 'Phòng chiếu'],
        ['start_time', 'DATETIME', 'NOT NULL', 'Giờ bắt đầu'],
        ['end_time', 'DATETIME', 'NOT NULL, CHECK > start_time', 'Giờ kết thúc'],
        ['base_price', 'DECIMAL(12,2)', 'NOT NULL, CHECK >= 0', 'Giá vé cơ bản'],
        ['status', 'VARCHAR(20)', "NOT NULL, DEFAULT 'SCHEDULED'", 'SCHEDULED | ONGOING | COMPLETED | CANCELLED'],
    ]
)

# show_seats table
add_heading_styled(new_doc, '6.6. Bảng show_seats', level=2)
add_formatted_table(new_doc,
    ['Cột', 'Kiểu', 'Ràng buộc', 'Mô tả'],
    [
        ['show_seat_id', 'VARCHAR(20)', 'PK', 'Mã ghế-suất chiếu'],
        ['showtime_id', 'VARCHAR(20)', 'FK → showtimes', 'Suất chiếu'],
        ['seat_id', 'VARCHAR(20)', 'FK → seats', 'Ghế'],
        ['price', 'DECIMAL(12,2)', 'NOT NULL, CHECK >= 0', 'Giá cho ghế này'],
        ['status', 'VARCHAR(20)', "NOT NULL, DEFAULT 'AVAILABLE'", 'AVAILABLE | HELD | BOOKED'],
        ['held_until', 'DATETIME', 'NULL', 'Hạn giữ chỗ'],
    ]
)
add_paragraph_styled(new_doc, 'UNIQUE constraint: (showtime_id, seat_id)', italic=True)

# invoices table
add_heading_styled(new_doc, '6.7. Bảng invoices', level=2)
add_formatted_table(new_doc,
    ['Cột', 'Kiểu', 'Ràng buộc', 'Mô tả'],
    [
        ['invoice_id', 'VARCHAR(20)', 'PK', 'Mã hóa đơn'],
        ['employee_id', 'VARCHAR(20)', 'FK → employees', 'Nhân viên tạo'],
        ['created_at', 'DATETIME', 'NOT NULL, DEFAULT GETDATE()', 'Thời gian tạo'],
        ['total_amount', 'DECIMAL(12,2)', 'NOT NULL, CHECK >= 0', 'Tổng tiền'],
        ['payment_method', 'VARCHAR(20)', "NOT NULL, CHECK IN ('CASH','CARD','TRANSFER')", 'Phương thức TT'],
        ['status', 'VARCHAR(20)', "NOT NULL, DEFAULT 'PENDING'", 'PENDING | CONFIRMED | CANCELLED'],
        ['note', 'NTEXT', 'NULL', 'Ghi chú'],
    ]
)

# tickets table
add_heading_styled(new_doc, '6.8. Bảng tickets', level=2)
add_formatted_table(new_doc,
    ['Cột', 'Kiểu', 'Ràng buộc', 'Mô tả'],
    [
        ['ticket_id', 'VARCHAR(20)', 'PK', 'Mã vé'],
        ['invoice_id', 'VARCHAR(20)', 'FK → invoices', 'Hóa đơn'],
        ['show_seat_id', 'VARCHAR(20)', 'FK → show_seats', 'Ghế suất chiếu'],
        ['discount_id', 'VARCHAR(20)', 'FK → discounts, NULL', 'Giảm giá (nếu có)'],
        ['barcode', 'VARCHAR(100)', 'UNIQUE', 'Mã vạch'],
        ['original_price', 'DECIMAL(12,2)', 'NOT NULL, CHECK >= 0', 'Giá gốc'],
        ['discount_amount', 'DECIMAL(12,2)', 'NOT NULL, DEFAULT 0', 'Số tiền giảm'],
        ['final_price', 'DECIMAL(12,2)', 'NOT NULL, CHECK >= 0', 'Giá cuối cùng'],
        ['issued_at', 'DATETIME', 'NOT NULL, DEFAULT GETDATE()', 'Thời gian xuất vé'],
        ['status', 'VARCHAR(20)', "NOT NULL, DEFAULT 'VALID'", 'VALID | USED | CANCELLED'],
    ]
)

# discounts table
add_heading_styled(new_doc, '6.9. Bảng discounts', level=2)
add_formatted_table(new_doc,
    ['Cột', 'Kiểu', 'Ràng buộc', 'Mô tả'],
    [
        ['discount_id', 'VARCHAR(20)', 'PK', 'Mã giảm giá'],
        ['discount_name', 'NVARCHAR(100)', 'NOT NULL', 'Tên chương trình'],
        ['discount_type', 'VARCHAR(20)', "NOT NULL, CHECK IN ('HOLIDAY','GROUP','SPECIAL')", 'Loại giảm giá'],
        ['discount_rate', 'DECIMAL(5,2)', 'NOT NULL, CHECK 0-1', 'Tỷ lệ giảm'],
        ['valid_from', 'DATE', 'NOT NULL', 'Ngày bắt đầu'],
        ['valid_to', 'DATE', 'NOT NULL', 'Ngày kết thúc'],
        ['min_ticket_quantity', 'INT', 'NOT NULL, DEFAULT 1', 'Số vé tối thiểu'],
        ['is_active', 'BIT', 'NOT NULL, DEFAULT 1', 'Đang hoạt động'],
        ['description', 'NTEXT', 'NULL', 'Mô tả'],
    ]
)

# ============================================================
# 7. SƠ ĐỒ ERD (MÔ TẢ VĂN BẢN)
# ============================================================
add_heading_styled(new_doc, '7. Sơ đồ ERD (Entity Relationship Diagram)', level=1)

add_paragraph_styled(new_doc, 'Mô tả quan hệ khóa ngoại:', bold=True)
erd_relations = [
    'seats.hall_id → cinema_halls.hall_id',
    'showtimes.movie_id → movies.movie_id',
    'showtimes.hall_id → cinema_halls.hall_id',
    'show_seats.showtime_id → showtimes.showtime_id',
    'show_seats.seat_id → seats.seat_id',
    'invoices.employee_id → employees.employee_id',
    'tickets.invoice_id → invoices.invoice_id',
    'tickets.show_seat_id → show_seats.show_seat_id',
    'tickets.discount_id → discounts.discount_id (nullable)',
]
for r in erd_relations:
    add_bullet(new_doc, r)

# ============================================================
# 8. ENUM TỔNG HỢP
# ============================================================
add_heading_styled(new_doc, '8. Enum tổng hợp', level=1)

add_formatted_table(new_doc,
    ['Enum', 'Giá trị', 'Dùng ở'],
    [
        ['EmployeeRole', 'MANAGER, STAFF', 'Employee.role'],
        ['EmployeeStatus', 'ACTIVE, INACTIVE', 'Employee.status'],
        ['MovieStatus', 'ACTIVE, INACTIVE', 'Movie.status'],
        ['SeatType', 'REGULAR, VIP, ACCESSIBLE', 'Seat.seatType'],
        ['ShowtimeStatus', 'SCHEDULED, ONGOING, COMPLETED, CANCELLED', 'Showtime.status'],
        ['ShowSeatStatus', 'AVAILABLE, HELD, BOOKED', 'ShowSeat.status'],
        ['PaymentMethod', 'CASH, CARD, TRANSFER', 'Invoice.paymentMethod'],
        ['InvoiceStatus', 'PENDING, CONFIRMED, CANCELLED', 'Invoice.status'],
        ['TicketStatus', 'VALID, USED, CANCELLED', 'Ticket.status'],
        ['DiscountType', 'HOLIDAY, GROUP, SPECIAL', 'Discount.discountType'],
    ],
    col_widths=[3.5, 6.5, 6]
)


# ─── Save ───
print(f"Saving updated document to {OUTPUT_PATH}...")
new_doc.save(OUTPUT_PATH)
print("✅ Done! Document has been updated successfully.")
print(f"📄 Backup saved at: {BACKUP_PATH}")
