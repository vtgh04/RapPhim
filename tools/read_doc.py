from docx import Document
import sys

doc = Document(r'c:\Users\ADMIN\Desktop\CinemaProject\RapPhim\docs\TaiLieu.docx')
with open(r'c:\Users\ADMIN\Desktop\CinemaProject\RapPhim\tools\tailieu_dump.txt', 'w', encoding='utf-8') as f:
    for i, p in enumerate(doc.paragraphs):
        f.write(f"[{i}][{p.style.name}] {p.text}\n")
    
    # Also dump tables
    for ti, table in enumerate(doc.tables):
        f.write(f"\n--- TABLE {ti} ---\n")
        for ri, row in enumerate(table.rows):
            cells = [cell.text for cell in row.cells]
            f.write(f"  Row {ri}: {' | '.join(cells)}\n")

print("Done")
